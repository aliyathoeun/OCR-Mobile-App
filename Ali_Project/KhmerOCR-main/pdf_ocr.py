import os
import pytesseract
import cv2
import numpy as np
from docx import Document
from docx.shared import Pt
from pdf2image import convert_from_path

# Define the PDF file input and output folder for docx files
pdf_file = 'earth.pdf'  # Path to the input PDF file
docx_output_folder = 'document_output'  # Folder to save the final .docx file

# Make sure the output directory exists
if not os.path.exists(docx_output_folder):
    os.makedirs(docx_output_folder)

# Extract the base name of the PDF file (without extension)
pdf_filename = os.path.splitext(os.path.basename(pdf_file))[0]

# Create a new Document for the combined content
doc = Document()

# Convert PDF pages to images
images = convert_from_path(pdf_file)

# Function to choose the best PSM mode based on the content
def select_psm_mode(image):
    """
    Select the appropriate PSM mode based on the content of the image.
    Returns the best Tesseract PSM config.
    """
    # Use edge detection to determine if the image contains tables (simple heuristic)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 50, 150)
    
    # If many edges are detected (likely a table), use PSM 4 or 6
    if np.sum(edges) > 5000:
        return r'-l khm+eng --psm 4'  # PSM 4 or 6 are better for tables
    else:
        return r'-l khm+eng --psm 3'  # PSM 3 for general text

# Loop over all pages (images) in the PDF
for page_number, image in enumerate(images, start=1):
    # Convert the PIL image to OpenCV format
    image_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
    
    # Preprocess the image (resize, blur, threshold)
    gray_image = cv2.cvtColor(image_cv, cv2.COLOR_BGR2GRAY)
    resized_image = cv2.resize(gray_image, (int(gray_image.shape[1] * 1.75), int(gray_image.shape[0] * 1.75)))
    blurred_image = cv2.GaussianBlur(resized_image, (5, 5), 0)
    _, threshold = cv2.threshold(blurred_image, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    
    # Select the best PSM mode based on content
    custom_config = select_psm_mode(image_cv)
    
    # OCR using Tesseract with the selected PSM mode
    text = pytesseract.image_to_string(threshold, config=custom_config)
    
    # Add the OCR text from this page into the document
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)  # Add the OCR text to the document
    
    # Change the font to Khmer OS Battambong
    run.font.name = 'Khmer OS Battambong'  # Set the font to Khmer OS Battambong
    run.font.size = Pt(13)  # Font size (you can adjust this as needed)
    
    # Add a page break after each page's text (optional)
    doc.add_paragraph()  # Adds a blank line to separate pages visually in the docx

    print(f"Processed page {page_number} and added to the combined document.")
    print("-" * 40)

# Save the combined document with the same name as the input PDF
output_docx_path = os.path.join(docx_output_folder, f'{pdf_filename}.docx')
doc.save(output_docx_path)

print(f"Combined document saved to {output_docx_path}")
