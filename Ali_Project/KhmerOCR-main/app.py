import os
import cv2
import pytesseract
import numpy as np
from flask import *
from docx import Document
from docx.shared import Pt, RGBColor
from pdf2image import convert_from_path
from werkzeug.utils import secure_filename
import io
from PIL import Image as PILImage
from joblib import load
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import time
import psutil
from flask_cors import CORS

app = Flask(__name__, static_folder='statics')
CORS(app)  # Enable CORS for all routes

# Configuration for directories
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'document_output')

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

# Ensure directories exist
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)

# Load font classification model and set up class labels
BASE_DIR = os.path.dirname(os.path.realpath(__file__))
MODEL_DIR = os.path.join(BASE_DIR, 'models')
model_path = os.path.join(MODEL_DIR, 'svm_model.joblib')
model = load(model_path)
class_labels = {
    0: "Khmer OS",
    1: "Khmer OS Battambong",
    2: "Khmer OS Siemreap",
    # Add more classes as needed
}

def preprocess_image(image, target_size=(64, 64)):
    """Preprocess the image for font classification."""
    image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    image = cv2.resize(image, target_size)
    image = image / 255.0  # Normalize to [0, 1]
    return image  # Removed expand_dims to handle flattening later

def detect_font_for_image(image):
    """Predict the font of a given image region."""
    first = first_word(image)  # Extract the first word region
    processed_image = preprocess_image(first)  # Preprocess the image
    
    flattened_image = processed_image.flatten().reshape(1, -1)  # Shape to (1, n_features)
    
    predictions = model.predict(flattened_image)  # SVM's predict doesn't need softmax
    predicted_class = predictions[0]  # SVM predict outputs directly the class
    

    return class_labels[predicted_class]

def set_font(run, font_name, font_size=12):
    """Set the font style for a run of text."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)  # Set the font color to black (default)

def first_word(image):
    """Extract the first word from the image by using OCR and cropping it."""
    # Convert image to grayscale for better OCR accuracy
    gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    # Use pytesseract to get OCR data
    text_data = pytesseract.image_to_data(gray_image, config=r'-l khm+eng --psm 6', output_type=pytesseract.Output.DICT)

    # Find the first word's bounding box
    first_word = ""
    for i, word in enumerate(text_data['text']):
        if word.strip():
            first_word = word.strip()  # Get the first word from OCR result
            x, y, w, h = (text_data['left'][i], text_data['top'][i], text_data['width'][i], text_data['height'][i])
            break  # Stop once we find the first word

    if first_word:
        # Crop the image to get the first word region
        cropped_image = image[y:y+h, x:x+w]
        return cropped_image
    else:
        return None, None
    
def detect_font_for_pdf(pdf_file):
    """Detect font for the first page of a PDF."""
    # Convert the first page of the PDF to an image
    images = convert_from_path(pdf_file)
    
    if not images:
        raise ValueError("Failed to convert PDF to image.")
    
    # Convert the first page to an OpenCV image (for further processing)
    image = images[0]  # Get the first page as a PIL Image
    image_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)  # Convert to OpenCV image
    
    font_name = detect_font_for_image(image_cv)
    
    return font_name


def perform_ocr_on_image(image, doc):
    # Convert the image to grayscale
    gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    
    # Perform OCR on the image
    text_data = pytesseract.image_to_data(
        gray_image, config=r'-l khm+eng --psm 6', output_type=pytesseract.Output.DICT
    )

    # Extract the full text from the OCR result
    result_text = ""
    for i in range(len(text_data['text'])):
        if int(text_data['conf'][i]) > 0:  # Filter by confidence (only use high-confidence words)
            word = text_data['text'][i].strip()
            if word:
                result_text += word + " "

    # Add the OCR result text to the document
    paragraph = doc.add_paragraph(result_text.strip())
    for run in paragraph.runs:
            run.font.name = "Khmer OS Battambang"

            # Ensure Word XML compatibility for Khmer font
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Khmer OS Battambang')

            # Set font size (optional)
            run.font.size = Pt(12)

    
def perform_ocr_on_pdf(pdf_file, doc):
    # Convert PDF to a list of PIL Image objects
    images = convert_from_path(pdf_file)

    # Iterate through the images (one per PDF page)
    for page_num, image in enumerate(images):
        # Convert the PIL image to OpenCV format
        image_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

        # Convert the image to grayscale
        gray_image = cv2.cvtColor(image_cv, cv2.COLOR_BGR2GRAY)

        # Perform OCR on the image
        text_data = pytesseract.image_to_data(
            gray_image, config=r'-l khm+eng --psm 6', output_type=pytesseract.Output.DICT
        )

        # Extract the full text from the OCR result
        result_text = ""
        for i in range(len(text_data['text'])):
            if int(text_data['conf'][i]) > 0:  # Filter by confidence (only use high-confidence words)
                word = text_data['text'][i].strip()
                if word:
                    result_text += word + " "

        # Add the OCR result text to the document (for each page)
        paragraph = doc.add_paragraph(result_text.strip())

        # Set the font to Khmer OS Battambang
        for run in paragraph.runs:
            run.font.name = "Khmer OS Battambang"

            # Ensure Word XML compatibility for Khmer font
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Khmer OS Battambang')

            # Set font size (optional)
            run.font.size = Pt(12)

        # Optionally, add page number or other details (e.g., "Page X")
        doc.add_paragraph



# Function to extract text from the docx
def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    doc_text = ""

    for para in doc.paragraphs:
        for run in para.runs:
            doc_text += run.text + " "  # Concatenate the runs to get the full paragraph text

    return doc_text


@app.route('/')
def index():
    error_message = request.args.get('error')  # Get error message if any
    return render_template('index.html', error_message=error_message)

@app.route('/detect', methods=['POST'])
def detect_font():
    if 'file' not in request.files:
        return redirect(url_for('index', error="No file uploaded."))

    file = request.files['file']
    if file.filename == '':
        return redirect(url_for('index', error="No file selected."))

    # Save uploaded file
    filename = secure_filename(file.filename)
    upload_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(upload_path)

    # Determine file extension and validate
    file_ext = os.path.splitext(filename)[1].lower()
    if file_ext not in ['.pdf', '.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
        return redirect(url_for('predict', error="Unsupported file type. Please upload a PDF or an image file."))

    # Detect font (no OCR here)
    try:
        if file_ext == '.pdf':
            # For PDFs, handle font detection (assuming a specific process for PDFs)
            font_name = detect_font_for_pdf(upload_path)
        else:
            # For images, handle font detection
            img = PILImage.open(upload_path)
            image_cv = np.array(img)
            if img.mode == 'RGB':
                image_cv = cv2.cvtColor(image_cv, cv2.COLOR_RGB2BGR)
            font_name = detect_font_for_image(image_cv)  # Adjust this function to your logic
    except Exception as e:
        return redirect(url_for('predict', error=f"Error detecting font: {str(e)}"))

    # Return font detection result
    return render_template('predict.html', font_name=font_name)


@app.route('/ocr', methods=['POST'])
def ocr_processing():
    if 'file' not in request.files:
        return redirect(url_for('ocr', error="No file uploaded."))

    file = request.files['file']
    if file.filename == '':
        return redirect(url_for('ocr', error="No file selected."))

    # Save uploaded file
    filename = secure_filename(file.filename)
    upload_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(upload_path)

    # Determine output file path
    output_filename = f"{os.path.splitext(filename)[0]}.docx"
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

    # Perform OCR (no font detection here)
    doc = Document()
    try:
        if os.path.splitext(filename)[1].lower() == '.pdf':
            # OCR for PDFs
            perform_ocr_on_pdf(upload_path, doc)
            doc.save(output_path)
        else:
            # OCR for images
            img = PILImage.open(upload_path)
            image_cv = np.array(img)
            if img.mode == 'RGB':
                image_cv = cv2.cvtColor(image_cv, cv2.COLOR_RGB2BGR)
            perform_ocr_on_image(image_cv, doc)  # Perform OCR
            doc.save(output_path)
    except Exception as e:
        return redirect(url_for('ocr', error=f"Error processing OCR: {str(e)}"))

    # Extract full text from the .docx file
    doc_text = extract_text_from_docx(output_path)

    # Return OCR result
    return render_template('ocr.html', doc_text=doc_text, download_link=output_filename)


@app.route('/download/<filename>')
def download(filename):
    return send_from_directory(app.config['OUTPUT_FOLDER'], filename, as_attachment=True)

import psutil

@app.before_request
def log_resource_usage():
    process = psutil.Process()  # Get the current process
    mem_info = process.memory_info()  # Get memory usage
    cpu_percent = process.cpu_percent(interval=0.1)  # Get CPU usage
    num_threads = process.num_threads()  # Get the number of threads
    open_files = process.open_files()  # Get the open files (optional)

    print("\n--- Resource Usage ---")
    print(f"CPU Usage: {cpu_percent:.2f}%")
    print(f"Memory Usage: {mem_info.rss / (1024 * 1024):.2f} MB")  # Convert bytes to MB
    print(f"Threads: {num_threads}")
    print(f"Open Files: {len(open_files)}")
    print("-----------------------")


if __name__ == "__main__":
    app.run(debug=True, host='0.0.0.0', port=5001)