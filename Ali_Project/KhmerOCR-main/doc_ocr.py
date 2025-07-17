import os
import pytesseract
import cv2
from docx import Document
from docx.shared import Pt

# Define the folder containing images and where the OCR results will be saved
image_folder = 'watermark'
docx_output_folder = 'document_output'  # Folder to save the final .docx files

# Make sure the output directory exists
if not os.path.exists(docx_output_folder):
    os.makedirs(docx_output_folder)

# Loop over all images in the image folder
image_files = [f for f in os.listdir(image_folder) if os.path.isfile(os.path.join(image_folder, f))]
for image_file in image_files:
    image_path = os.path.join(image_folder, image_file)
    
    # Read and process the image
    image = cv2.imread(image_path)
    gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    
    # Resize the image (175% of original size)
    height, width = gray_image.shape
    new_width = int(width * 1.75)
    new_height = int(height * 1.75)
    resized_image = cv2.resize(gray_image, (new_width, new_height))
    
    # Binarize the image using Otsu's method
    _, threshold = cv2.threshold(resized_image, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    
    # OCR using Tesseract (Khmer and English)
    custom_config = r'-l khm+eng --psm 3'
    text = pytesseract.image_to_string(threshold, config=custom_config)
    
    # Create a new Document for each image
    doc = Document()
    
    # Add a paragraph and change the font style
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)  # Add the OCR text to the document
    
    # Change the font to Khmer OS Battambong
    run.font.name = 'Khmer OS Battambong'  # Set the font to Khmer OS Battambong
    run.font.size = Pt(13)  # Font size (you can adjust this as needed)
    
    # Define the path for the output .docx file
    output_docx_path = os.path.join(docx_output_folder, f'{os.path.splitext(image_file)[0]}.docx')
    
    # Save the document as a separate .docx file
    doc.save(output_docx_path)
    
    # Print the text that was extracted and the path of the saved file
    print(f"Text from {image_file} saved to {output_docx_path}")
    print("-" * 40)
